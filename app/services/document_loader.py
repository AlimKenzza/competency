"""
document_loader.py
------------------
Модуль загрузки и хранения данных (раздел 7.1 ТЗ).
Извлекает текст из документов форматов TXT, DOCX, PDF, CSV, XLSX.

Возвращает либо «сырой текст» (для свободных документов, идущих в извлечение
сущностей), либо структурированный DataFrame (для CSV/XLSX).
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

import pandas as pd

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf", ".csv", ".xlsx", ".xls"}


def detect_format(path: str | Path) -> str:
    ext = Path(path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Неподдерживаемый формат: {ext}. "
                         f"Поддерживаются: {sorted(SUPPORTED_EXTENSIONS)}")
    return ext


def load_text(path: str | Path) -> str:
    """Извлекает текстовое содержимое из TXT/DOCX/PDF."""
    ext = detect_format(path)
    path = Path(path)

    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if ext == ".docx":
        return _load_docx(path)

    if ext == ".pdf":
        return _load_pdf(path)

    raise ValueError(f"Формат {ext} не является текстовым. "
                     f"Используйте load_table() для CSV/XLSX.")


def load_table(path: str | Path) -> pd.DataFrame:
    """Загружает табличные данные из CSV/XLSX."""
    ext = detect_format(path)
    path = Path(path)
    if ext == ".csv":
        return pd.read_csv(path)
    if ext in (".xlsx", ".xls"):
        return pd.read_excel(path)
    raise ValueError(f"Формат {ext} не является таблицей.")


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Установите python-docx: pip install python-docx")
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Установите pypdf: pip install pypdf")
    reader = PdfReader(str(path))
    texts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            texts.append(t)
    return "\n".join(texts)


def load_any(path: str | Path) -> dict:
    """
    Универсальная загрузка. Возвращает dict:
      {"kind": "text"|"table", "text": str|None, "table": DataFrame|None, "format": ext}
    """
    ext = detect_format(path)
    if ext in (".csv", ".xlsx", ".xls"):
        return {"kind": "table", "text": None, "table": load_table(path), "format": ext}
    else:
        return {"kind": "text", "text": load_text(path), "table": None, "format": ext}
