"""
report_service.py
-----------------
Модуль отчётности (раздел 7.7 ТЗ). Формирует:
  - карточку профессионального профиля;
  - отчёт по сопоставлению компетенций;
  - сравнительную таблицу по источникам;
  - аналитическую сводку по качеству.
Экспорт в форматы PDF, DOCX, XLSX (раздел 7.7).
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from app.core.config import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def export_profile_xlsx(profile, items, out_path: str | None = None) -> str:
    """Экспорт профиля в XLSX с несколькими листами."""
    import pandas as pd

    if out_path is None:
        out_path = str(Path(settings.report_dir) /
                       f"profile_{profile.id}_{_ts()}.xlsx")

    # Лист 1: компетенции профиля
    rows = []
    for it in items:
        rows.append({
            "Компетенция": it.title,
            "Тип": it.competency_type,
            "Уровень владения": it.proficiency_level,
            "Релевантность": round(it.similarity, 4),
            "Источников": it.n_sources,
            "Подтверждено экспертом": "Да" if it.expert_confirmed else "Нет",
        })
    df_items = pd.DataFrame(rows)

    # Лист 2: метрики качества
    qm = profile.quality_metrics or {}
    df_quality = pd.DataFrame([
        {"Показатель": "Полнота", "Значение": qm.get("completeness", "")},
        {"Показатель": "Непротиворечивость", "Значение": qm.get("consistency", "")},
        {"Показатель": "Применимость", "Значение": qm.get("applicability", "")},
        {"Показатель": "Согласованность источников", "Значение": qm.get("source_agreement", "")},
        {"Показатель": "Число компетенций", "Значение": qm.get("n_items", "")},
    ])

    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        df_items.to_excel(writer, sheet_name="Компетенции", index=False)
        df_quality.to_excel(writer, sheet_name="Качество", index=False)

    logger.info(f"XLSX-отчёт сохранён: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def export_profile_docx(profile, items, out_path: str | None = None) -> str:
    """Экспорт карточки профиля в DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if out_path is None:
        out_path = str(Path(settings.report_dir) /
                       f"profile_{profile.id}_{_ts()}.docx")

    doc = Document()

    # Заголовок
    h = doc.add_heading("Карточка профессионального профиля", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(profile.role_title, level=1)
    if profile.role_description:
        doc.add_paragraph(profile.role_description)

    p = doc.add_paragraph()
    p.add_run(f"Дата формирования: {profile.created_at:%Y-%m-%d %H:%M}").italic = True

    # Метрики качества
    doc.add_heading("Показатели качества профиля", level=2)
    qm = profile.quality_metrics or {}
    quality_table = doc.add_table(rows=1, cols=2)
    quality_table.style = "Light Grid Accent 1"
    quality_table.rows[0].cells[0].text = "Показатель"
    quality_table.rows[0].cells[1].text = "Значение"
    for label, key in [
        ("Полнота", "completeness"),
        ("Непротиворечивость", "consistency"),
        ("Применимость", "applicability"),
        ("Согласованность источников", "source_agreement"),
    ]:
        row = quality_table.add_row().cells
        row[0].text = label
        row[1].text = str(qm.get(key, "—"))

    # Таблица компетenций
    doc.add_heading("Компетенции профиля", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Компетенция"
    hdr[2].text = "Тип"
    hdr[3].text = "Уровень"
    hdr[4].text = "Релевантность"
    hdr[5].text = "Эксперт"
    for i, it in enumerate(items, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = it.title
        row[2].text = it.competency_type
        row[3].text = it.proficiency_level
        row[4].text = f"{it.similarity:.3f}"
        row[5].text = "✓" if it.expert_confirmed else "—"

    doc.save(out_path)
    logger.info(f"DOCX-отчёт сохранён: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def export_profile_pdf(profile, items, out_path: str | None = None) -> str:
    """Экспорт карточки профиля в PDF (reportlab)."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if out_path is None:
        out_path = str(Path(settings.report_dir) /
                       f"profile_{profile.id}_{_ts()}.pdf")

    # Регистрируем шрифт с кириллицей
    font_name = _register_cyrillic_font()

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCyr", parent=styles["Title"], fontName=font_name)
    h2_style = ParagraphStyle(
        "H2Cyr", parent=styles["Heading2"], fontName=font_name)
    normal_style = ParagraphStyle(
        "NormalCyr", parent=styles["Normal"], fontName=font_name, fontSize=9)

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm)
    flow = []
    flow.append(Paragraph("Карточка профессионального профиля", title_style))
    flow.append(Spacer(1, 0.3*cm))
    flow.append(Paragraph(profile.role_title, h2_style))
    if profile.role_description:
        flow.append(Paragraph(profile.role_description, normal_style))
    flow.append(Spacer(1, 0.3*cm))

    # Метрики
    qm = profile.quality_metrics or {}
    flow.append(Paragraph("Показатели качества", h2_style))
    q_data = [["Показатель", "Значение"]]
    for label, key in [
        ("Полнота", "completeness"),
        ("Непротиворечивость", "consistency"),
        ("Применимость", "applicability"),
        ("Согласованность источников", "source_agreement"),
    ]:
        q_data.append([label, str(qm.get(key, "—"))])
    q_table = Table(q_data, colWidths=[8*cm, 4*cm])
    q_table.setStyle(_table_style(font_name))
    flow.append(q_table)
    flow.append(Spacer(1, 0.4*cm))

    # Компетенции
    flow.append(Paragraph("Компетенции профиля", h2_style))
    data = [["№", "Компетенция", "Тип", "Уровень", "Релев.", "Эксперт"]]
    for i, it in enumerate(items, 1):
        data.append([
            str(i),
            Paragraph(it.title, normal_style),
            it.competency_type,
            it.proficiency_level,
            f"{it.similarity:.3f}",
            "✓" if it.expert_confirmed else "—",
        ])
    table = Table(data, colWidths=[1*cm, 7*cm, 2.5*cm, 3*cm, 1.8*cm, 1.7*cm])
    table.setStyle(_table_style(font_name))
    flow.append(table)

    doc.build(flow)
    logger.info(f"PDF-отчёт сохранён: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# Вспомогательные
# ---------------------------------------------------------------------------

def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _register_cyrillic_font() -> str:
    """Регистрирует шрифт с поддержкой кириллицы для reportlab."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/Library/Fonts/Arial.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont("CyrFont", path))
                return "CyrFont"
            except Exception:
                continue
    # Фолбэк на встроенный (может не поддерживать кириллицу)
    return "Helvetica"


def _table_style(font_name: str):
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E75B6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
    ])
